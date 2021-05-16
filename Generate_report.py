import os
import glob
import shutil
from docx import Document
from docx.shared import Inches

document = Document()
date = input('Input date in YYYY-MM-DD: ') #date of report to be printed as input
img_dir = "NOTSAFE/{}".format(date) # Enter Directory of all images
document.add_heading('NOTSAFE on {}'.format(date), 0) #heading of document file
document.add_paragraph('This report is generated as a part of Masters Thesis in Construction technology and Management course at IIT Bombay by Sankarlal.')

os.chdir(img_dir)
for file in glob.glob("*.jpg"):
    document.add_heading(file, level=1)     #heading of photo
    document.add_picture(file, width=Inches(4))
    #document.add_page_break()

document.save("/home/ctam/Desktop/SankarMTP/stage2/yolov3/Reports/{}.docx".format(date)) #saving document to reports folder
print('Report of {} saved to Reports folder'.format(date))
